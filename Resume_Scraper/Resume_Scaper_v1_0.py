import fitz  # PyMuPDF
from docx import Document

def extract_keywords_from_pdf(pdf_path, keywords):
    doc = fitz.open(pdf_path)
    text = ''
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()

    # Check for keywords in the extracted text
    found_keywords = [word for word in keywords if word.lower() in text.lower()]

    return found_keywords

def create_word_document(output_path, keywords_found):
    doc = Document()
    doc.add_heading('Keywords Found', level=1)

    for keyword in keywords_found:
        doc.add_paragraph(keyword)

    doc.save(output_path)

if __name__ == "__main__":
    # PDF paths
    pdf1_path = 'C:\\Users\\kmoye\\OneDrive\\7. School docs\\Univerity Transcripts'
    pdf2_path = 'C:\\Users\\kmoye\\OneDrive\\7. School docs\\Univerity Transcripts'

    # Keywords to search for
    search_keywords = ['example', 'keyword', 'python', 'document']

    # Extract keywords from both PDFs
    keywords_pdf1 = extract_keywords_from_pdf(pdf1_path, search_keywords)
    keywords_pdf2 = extract_keywords_from_pdf(pdf2_path, search_keywords)

    # Combine keywords from both PDFs
    all_keywords_found = set(keywords_pdf1 + keywords_pdf2)

    # Create Word document with found keywords
    create_word_document('output_keywords.docx', all_keywords_found)
